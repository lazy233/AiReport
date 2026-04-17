"""Word 文档解析：提取标题、段落、表格为结构化 sections。"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as _DocType
from docx.table import Table
from docx.text.paragraph import Paragraph


def set_table_cell_text_preserve_style(cell: Any, text: str) -> None:
    """
    写入表格单元格文本并尽量保留模板中的字体与段落格式。

    直接使用 cell.text = ... 会清空单元格并按文档默认样式重建段落，导致原模板字体（含中文 East Asia）
    与加粗等丢失；这里在改写前复制首个 rPr / 首段 pPr 的 OOXML，写回后再套用到新 run/段上。
    """
    rpr_template = None
    ppr_template = None
    if cell.paragraphs:
        pp0 = cell.paragraphs[0]._element.pPr
        if pp0 is not None:
            ppr_template = deepcopy(pp0)
    for para in cell.paragraphs:
        for run in para.runs:
            rp = run._element.rPr
            if rp is not None:
                rpr_template = deepcopy(rp)
                break
        if rpr_template is not None:
            break

    cell.text = text

    if ppr_template is not None:
        for para in cell.paragraphs:
            pel = para._element
            old = pel.pPr
            if old is not None:
                pel.remove(old)
            pel.insert(0, deepcopy(ppr_template))

    if rpr_template is not None:
        for para in cell.paragraphs:
            for run in para.runs:
                rel = run._element
                old = rel.rPr
                if old is not None:
                    rel.remove(old)
                rel.insert(0, deepcopy(rpr_template))


def _iter_block_items(doc: _DocType):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


def _heading_level_from_style_name(style_name: str) -> int:
    raw = (style_name or "").strip().lower()
    if not raw:
        return 0
    # 常见英文/中文样式名：Heading 1、标题 1
    for token in raw.replace("heading", "").replace("标题", "").split():
        if token.isdigit():
            lv = int(token)
            return min(9, max(1, lv))
    digits = "".join(ch for ch in raw if ch.isdigit())
    if digits:
        lv = int(digits)
        return min(9, max(1, lv))
    return 0


def parse_docx(docx_path: str | Path) -> dict[str, Any]:
    path = Path(docx_path)
    if not path.is_file():
        raise RuntimeError("Word 文件不存在，无法解析。")
    doc = Document(str(path))
    sections: list[dict[str, Any]] = []
    para_index = 0
    table_index = 0

    for blk in _iter_block_items(doc):
        if isinstance(blk, Paragraph):
            text = (blk.text or "").strip()
            if not text:
                continue
            para_index += 1
            style_name = blk.style.name if blk.style is not None else ""
            level = _heading_level_from_style_name(style_name)
            row: dict[str, Any] = {
                "index": len(sections) + 1,
                "source": "paragraph",
                "paragraph_index": para_index,
                "text": text,
                "style": style_name or "",
            }
            if level > 0:
                row["type"] = "heading"
                row["level"] = level
            else:
                row["type"] = "paragraph"
            sections.append(row)
            continue

        if isinstance(blk, Table):
            table_index += 1
            rows: list[list[str]] = []
            max_cols = 0
            for r in blk.rows:
                vals = [(c.text or "").strip() for c in r.cells]
                max_cols = max(max_cols, len(vals))
                rows.append(vals)
            sections.append(
                {
                    "index": len(sections) + 1,
                    "type": "table",
                    "source": "table",
                    "table_index": table_index,
                    "row_count": len(rows),
                    "col_count": max_cols,
                    "rows": rows,
                },
            )

    return {
        "section_count": len(sections),
        "sections": sections,
    }


"""Word 生成：根据学生数据调用大模型，决策并回填表格单元格内容。"""
from __future__ import annotations

import json
import os
import re
from pathlib import Path
from typing import Any

from docx import Document
import requests

from ppt_report import config
from ppt_report.models import db
from ppt_report.services.chapter_ref_images import is_safe_task_id
from ppt_report.services.llm_json import extract_json_from_text
from ppt_report.services.word_document import set_table_cell_text_preserve_style

_PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_.\-]+)\s*\}\}")
# 标签列可较长（如「平均每节课时长」「Classin 沟通频次」）
_KV_LINE_RE = re.compile(r"^\s*([^:：]{1,80})\s*[:：]\s*(.*?)\s*$")
_LABEL_ONLY_COLON_RE = re.compile(r"^\s*([^:：]{1,80})\s*[:：]\s*$")
# 宽松：标签内可含冒号以外的任意首段，值段至少 1 字（避免把「对接顾问/BD：」判错）
_KV_RELAXED_RE = re.compile(r"^(.+?)[:：]\s*(.+)\s*$")


def _normalize_key(s: str) -> str:
    return re.sub(r"[\s_:\-：]", "", str(s or "").strip().lower())


def _first_nonempty(d: dict[str, Any], *keys: str) -> str:
    for k in keys:
        v = str(d.get(k) or "").strip()
        if v:
            return v
    return ""


def _hour_cell(hours: dict[str, Any], ts: dict[str, Any], camel: str, snake: str) -> str:
    v = str(hours.get(camel) or "").strip()
    if v:
        return v
    return str(ts.get(snake) or "").strip()


def _build_student_value_map(record: dict[str, Any]) -> tuple[dict[str, str], dict[str, str]]:
    profile = record.get("profile") if isinstance(record.get("profile"), dict) else {}
    basic = profile.get("basic") if isinstance(profile.get("basic"), dict) else {}
    learning = profile.get("learning") if isinstance(profile.get("learning"), dict) else {}
    hours = profile.get("hours") if isinstance(profile.get("hours"), dict) else {}
    guidance = profile.get("guidance") if isinstance(profile.get("guidance"), dict) else {}
    term_summary = profile.get("term_summary") if isinstance(profile.get("term_summary"), dict) else {}
    raw: dict[str, str] = {
        "studentName": str(record.get("name") or basic.get("studentName") or "").strip(),
        "studentId": str(record.get("studentId") or basic.get("studentId") or "").strip(),
        "schoolName": str(basic.get("schoolName") or basic.get("school") or "").strip(),
        "grade": str(basic.get("grade") or basic.get("gradeLevel") or "").strip(),
        "semesterName": str(basic.get("semesterName") or basic.get("currentTerm") or "").strip(),
        "teacherName": str(basic.get("teacherName") or basic.get("plannerTeacher") or "").strip(),
        "plannerTeacher": str(basic.get("plannerTeacher") or "").strip(),
        "advisorTeacher": str(basic.get("advisorTeacher") or "").strip(),
        "gradeIntake": str(basic.get("gradeIntake") or "").strip(),
        "className": str(basic.get("className") or "").strip(),
        "product": str(basic.get("product") or "").strip(),
        "major": _first_nonempty(basic, "major", "majorName")
        or str(learning.get("majorDirection") or learning.get("goalMajor") or learning.get("major") or "").strip(),
        "majorDirection": str(learning.get("majorDirection") or "").strip(),
        "transferDate": _first_nonempty(
            basic,
            "transferDate",
            "caseTransferDate",
            "recordTransferDate",
            "switchCaseDate",
        )
        or str(basic.get("serviceStart") or basic.get("serviceStartDate") or "").strip(),
        # 与转案日期同源，供别名「对接日期」等推断
        "serviceStart": str(basic.get("serviceStart") or basic.get("serviceStartDate") or "").strip(),
        "goalSchool": str(learning.get("goalSchool") or "").strip(),
        "goalMajor": str(learning.get("goalMajor") or "").strip(),
        "strength_subjects": _first_nonempty(
            learning, "strength_subjects", "strongSubjects",
        ),
        "scores": _first_nonempty(learning, "scores", "intlScores"),
        "learning_good": _first_nonempty(learning, "learning_good", "learningStyle"),
        "learning_weak": _first_nonempty(learning, "learning_weak", "weakAreas"),
        "interests": _first_nonempty(learning, "interests", "interestSubjects"),
        "study_goal": _first_nonempty(learning, "study_goal", "studyIntent"),
        "career_goal": _first_nonempty(learning, "career_goal", "careerIntent"),
        "long_goal": _first_nonempty(learning, "long_goal", "longTermPlan"),
        "degree": _first_nonempty(learning, "degree"),
        "duration": _first_nonempty(learning, "duration"),
        "credits": _first_nonempty(learning, "credits"),
        "course_rule": _first_nonempty(learning, "course_rule"),
        "gpa_rule": _first_nonempty(learning, "gpa_rule"),
        "selection_rule": _first_nonempty(learning, "selection_rule"),
        "recommended_courses": _first_nonempty(learning, "recommended_courses"),
        "course_notes": _first_nonempty(learning, "course_notes"),
        "term_plan": _first_nonempty(learning, "term_plan"),
        "future_plan": _first_nonempty(learning, "future_plan"),
        "prep_courses": _first_nonempty(hours, "prep_courses", "previewSubjects"),
        "tutoring_courses": _first_nonempty(hours, "tutoring_courses", "tutorSubjects"),
        "totalHours": _hour_cell(hours, term_summary, "totalHours", "total_hours"),
        "usedHours": _hour_cell(hours, term_summary, "usedHours", "used_hours"),
        "remainingHours": _hour_cell(hours, term_summary, "remainingHours", "left_hours"),
        "total_hours": _hour_cell(hours, term_summary, "totalHours", "total_hours"),
        "used_hours": _hour_cell(hours, term_summary, "usedHours", "used_hours"),
        "left_hours": _hour_cell(hours, term_summary, "remainingHours", "left_hours"),
        "student_summary": str(term_summary.get("student_summary") or "").strip(),
        "school_ddl": str(term_summary.get("school_ddl") or "").strip(),
        "first_class_time": str(term_summary.get("first_class_time") or "").strip(),
        "first_class_note": str(term_summary.get("first_class_note") or "").strip(),
        "summer_work": str(term_summary.get("summer_work") or "").strip(),
        "term_work": str(term_summary.get("term_work") or "").strip(),
        "recorded_courses": str(term_summary.get("recorded_courses") or "").strip(),
        "grades": str(term_summary.get("grades") or "").strip(),
        "gpa": str(term_summary.get("gpa") or "").strip(),
        "target_gpa": str(term_summary.get("target_gpa") or "").strip(),
        "final_score": str(term_summary.get("final_score") or "").strip(),
        "services": str(term_summary.get("services") or "").strip(),
        "service_count": str(term_summary.get("service_count") or "").strip(),
        "class_count": str(term_summary.get("class_count") or "").strip(),
        "total_duration": str(term_summary.get("total_duration") or "").strip(),
        "avg_duration": str(term_summary.get("avg_duration") or "").strip(),
        "communication": str(term_summary.get("communication") or "").strip(),
        "next_goal": str(term_summary.get("next_goal") or "").strip(),
        "risk_courses": str(term_summary.get("risk_courses") or "").strip(),
        "suggestions": str(term_summary.get("suggestions") or "").strip(),
        "remarks": str(term_summary.get("remarks") or "").strip(),
        "termSummary": str(guidance.get("termSummary") or "").strip(),
        "courseFeedback": str(guidance.get("courseFeedback") or "").strip(),
        "shortTermAdvice": str(guidance.get("shortTermAdvice") or "").strip(),
        "longTermDevelopment": str(guidance.get("longTermDevelopment") or "").strip(),
        "content": str(record.get("content") or "").strip(),
        "remark": str(record.get("remark") or "").strip(),
    }
    values = {k: v for k, v in raw.items() if v}
    aliases: dict[str, str] = {}
    alias_pairs = {
        "姓名": "studentName",
        "学生姓名": "studentName",
        "学号": "studentId",
        "学校": "schoolName",
        "年级": "grade",
        "学期": "semesterName",
        "老师": "teacherName",
        "规划老师": "plannerTeacher",
        "教服老师": "advisorTeacher",
        "负责教服老师": "advisorTeacher",
        "专业": "major",
        "转案日期": "transferDate",
        "委托产品": "product",
        "班级": "className",
        "产品": "product",
        "对接顾问": "plannerTeacher",
        "对接顾问/BD": "plannerTeacher",
        "Classin沟通频次": "communication",
        "Classin 沟通频次": "communication",
        "包课课时总统计": "totalHours",
        "年级/入学时间": "gradeIntake",
        "对接日期": "transferDate",
        "服务开始日期": "serviceStart",
        "擅长科目": "strength_subjects",
        "语言/国际成绩": "scores",
        "擅长学习形式": "learning_good",
        "不擅长学习形式": "learning_weak",
        "兴趣方向": "interests",
        "升学意向": "study_goal",
        "就业意向": "career_goal",
        "长远目标": "long_goal",
        "学位": "degree",
        "学制": "duration",
        "学分要求": "credits",
        "课程要求": "course_rule",
        "GPA要求": "gpa_rule",
        "选课规则": "selection_rule",
        "推荐课程": "recommended_courses",
        "课程说明": "course_notes",
        "学期规划": "term_plan",
        "后续规划": "future_plan",
        "预习课程": "prep_courses",
        "同步辅导课程": "tutoring_courses",
        "目标学校": "goalSchool",
        "目标专业": "goalMajor",
        "规划方向": "majorDirection",
        "总课时": "totalHours",
        "已用课时": "usedHours",
        "剩余课时": "remainingHours",
        "学生情况": "student_summary",
        "校方DDL": "school_ddl",
        "首课时间": "first_class_time",
        "首课记录": "first_class_note",
        "暑期辅导": "summer_work",
        "学期辅导": "term_work",
        "录播完成": "recorded_courses",
        "成绩明细": "grades",
        "GPA": "gpa",
        "目标GPA": "target_gpa",
        "最终成绩": "final_score",
        "服务内容": "services",
        "服务次数": "service_count",
        "课程次数": "class_count",
        "总时长": "total_duration",
        "平均课时": "avg_duration",
        "沟通频次": "communication",
        "下阶段目标": "next_goal",
        "风险科目": "risk_courses",
        "下阶段建议": "suggestions",
        "学期总结备注": "remarks",
        "学期表现概述": "termSummary",
        "学期总结": "termSummary",
        "课程反馈": "courseFeedback",
        "短期建议": "shortTermAdvice",
        "长期发展": "longTermDevelopment",
        "备注": "remark",
        "内容": "content",
    }
    for k, key_name in alias_pairs.items():
        val = values.get(key_name, "")
        if val:
            aliases[_normalize_key(k)] = key_name
    return values, aliases


def _replace_placeholders(text: str, values: dict[str, str]) -> str:
    """将 {{fieldName}} 替换为 value_map 中的值。"""

    def _sub(m: re.Match[str]) -> str:
        key = (m.group(1) or "").strip()
        return str(values.get(key, "") or "")

    return _PLACEHOLDER_RE.sub(_sub, text or "")


def _replace_kv_lines(text: str, values: dict[str, str], aliases: dict[str, str]) -> str:
    """逐行匹配「标签: 值」；若标签能映射到 value_map 且该字段有值，则替换冒号后的内容。"""
    lines = (text or "").split("\n")
    out: list[str] = []
    for line in lines:
        m = _KV_LINE_RE.match(line)
        if not m:
            out.append(line)
            continue
        label = m.group(1) or ""
        nk = _normalize_key(label)
        key_name = aliases.get(nk)
        if not key_name:
            out.append(line)
            continue
        new_val = str(values.get(key_name, "") or "").strip()
        if not new_val:
            out.append(line)
            continue
        sep = "：" if "：" in line else ":"
        out.append(f"{label.strip()}{sep}{new_val}")
    return "\n".join(out)


def _deterministic_cell_text(text: str, values: dict[str, str], aliases: dict[str, str]) -> str:
    t = _replace_placeholders(text or "", values)
    return _replace_kv_lines(t, values, aliases)


# 模板左格仅文字、右格为 / 或空时常用行（无冒号版本）
_TWO_COL_LEFT_LABEL_KEYS: tuple[tuple[str, str], ...] = (
    ("包课课时总统计", "totalHours"),
    ("课程次数", "class_count"),
    ("课程总时长", "total_duration"),
    ("平均每节课时长", "avg_duration"),
)


def _plain_two_col_map() -> dict[str, str]:
    return {_normalize_key(lab): vkey for lab, vkey in _TWO_COL_LEFT_LABEL_KEYS}


def _right_is_placeholder(s: str) -> bool:
    t = (s or "").replace("\u3000", " ").replace("\xa0", " ").strip()
    if not t:
        return True
    return t in ("/", "／", "—", "-", "...", "…", "–", "－")


def _cell_suggest_llm_infer(text: str) -> bool:
    """空单元格、整格仅为占位符、或仅有「标签：」无值时，提示大模型结合档案推理填写。"""
    s = (text or "").strip()
    if not s:
        return True
    if _right_is_placeholder(s):
        return True
    fl = (s.split("\n")[0] or "").strip()
    if _LABEL_ONLY_COLON_RE.match(fl):
        return True
    return False


def _match_kv_first_line(first_line: str) -> tuple[str, str] | None:
    """解析首行「标签：值」，先严格再宽松。"""
    s = (first_line or "").strip()
    if not s:
        return None
    m = _KV_LINE_RE.match(s)
    if m:
        return (m.group(1).strip(), (m.group(2) or "").strip())
    m2 = _KV_RELAXED_RE.match(s)
    if m2:
        return (m2.group(1).strip(), (m2.group(2) or "").strip())
    return None


def _fix_two_column_rows(doc: Document, values: dict[str, str], aliases: dict[str, str]) -> int:
    """
    将误写在左格的「标签：值」拆成左「标签：」+ 右「值」；
    左格仅有「标签：」时把 value_map 填到右格；
    左格为无冒号标题且右格为占位时，按别名/包课行映射填右格。

    使用 table.cell(row,col) 按网格取格；合并单元格时两格可能为同一 _tc，无法在物理上拆成两格，
    此时改为左格内「标签：\\n值」排版，避免与「误跳过」导致永不修复。
    """
    plain = _plain_two_col_map()
    changed = 0
    for table in doc.tables:
        nrows = len(table.rows)
        ncols = len(table.columns)
        if ncols < 2:
            continue
        for ri in range(nrows):
            try:
                left = table.cell(ri, 0)
                right = table.cell(ri, 1)
            except (IndexError, ValueError):
                continue
            merged_same = left._tc is right._tc
            lt_raw = left.text or ""
            rt_raw = right.text or ""
            first_line = (lt_raw.split("\n")[0] or "").strip()

            kv = _match_kv_first_line(first_line)
            if kv:
                lab, val = kv[0], kv[1]
                if val and _right_is_placeholder(rt_raw):
                    sep = "：" if "：" in first_line else ":"
                    lines = lt_raw.split("\n")
                    lines[0] = f"{lab}{sep}"
                    body = "\n".join(lines).strip()
                    if merged_same:
                        extra = "\n".join(lines[1:]) if len(lines) > 1 else ""
                        pack = f"{lab}{sep}\n{val}"
                        if extra:
                            pack += "\n" + extra
                        set_table_cell_text_preserve_style(left, pack)
                    else:
                        set_table_cell_text_preserve_style(left, body)
                        set_table_cell_text_preserve_style(right, val)
                    changed += 1
                    continue

            m2 = _LABEL_ONLY_COLON_RE.match(first_line)
            if m2 and _right_is_placeholder(rt_raw):
                nk = _normalize_key(m2.group(1))
                key_name = aliases.get(nk)
                if key_name:
                    nv = str(values.get(key_name, "") or "").strip()
                    if nv:
                        if merged_same:
                            set_table_cell_text_preserve_style(left, f"{first_line.strip()}\n{nv}")
                        else:
                            set_table_cell_text_preserve_style(right, nv)
                        changed += 1
                        continue

            if (
                first_line
                and ":" not in first_line
                and "：" not in first_line
                and _right_is_placeholder(rt_raw)
            ):
                nk2 = _normalize_key(first_line)
                vk = plain.get(nk2)
                if vk:
                    nv = str(values.get(vk, "") or "").strip()
                    if nv:
                        if merged_same:
                            set_table_cell_text_preserve_style(
                                left, f"{first_line.strip()}\n{nv}",
                            )
                        else:
                            set_table_cell_text_preserve_style(right, nv)
                        changed += 1
                        continue
                key_name2 = aliases.get(nk2)
                if key_name2:
                    nv2 = str(values.get(key_name2, "") or "").strip()
                    if nv2:
                        if merged_same:
                            set_table_cell_text_preserve_style(
                                left, f"{first_line.strip()}\n{nv2}",
                            )
                        else:
                            set_table_cell_text_preserve_style(right, nv2)
                        changed += 1
    return changed


def _fix_two_column_rows_until_stable(
    doc: Document,
    values: dict[str, str],
    aliases: dict[str, str],
    *,
    max_rounds: int = 8,
) -> int:
    """合并格写入后可能影响相邻行，多轮直到无变化。"""
    total = 0
    for _ in range(max_rounds):
        n = _fix_two_column_rows(doc, values, aliases)
        total += n
        if n == 0:
            break
    return total


def _should_allow_llm_cell_update(original: str) -> bool:
    """
    为 True 时允许用大模型返回覆盖整格原文。
    空单元格、整格仅为 / 等占位、或仅有「标签：」无值时也必须为 True，否则模型无法推理填入内容。
    无冒号、无占位符的短固定文案（如单独「李老师」）为 False，避免误覆盖。
    """
    o = original or ""
    st = o.strip()
    if not st:
        return True
    if _right_is_placeholder(st):
        return True
    first = o.split("\n", 1)[0].strip()
    if _LABEL_ONLY_COLON_RE.match(first):
        return True
    if "{{" in o:
        return True
    if _KV_LINE_RE.match(first) or _match_kv_first_line(first):
        return True
    if len(o) > 320 or o.count("\n") >= 2:
        return True
    return False


def _extract_table_cells(doc: Document) -> tuple[list[dict[str, Any]], int]:
    cells: list[dict[str, Any]] = []
    table_count = 0
    for t_idx, table in enumerate(doc.tables):
        table_count += 1
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                otxt = (cell.text or "").strip()
                cells.append(
                    {
                        "cell_id": f"t{t_idx}_r{r_idx}_c{c_idx}",
                        "table_index": t_idx,
                        "row_index": r_idx,
                        "col_index": c_idx,
                        "original_text": otxt,
                        "suggest_infer": _cell_suggest_llm_infer(otxt),
                    },
                )
    return cells, table_count


def _llm_fill_table_cells(
    *,
    student_record: dict[str, Any],
    values: dict[str, str],
    aliases: dict[str, str],
    cells: list[dict[str, Any]],
) -> dict[str, str]:
    api_key = os.getenv("DASHSCOPE_API_KEY") or os.getenv("OPENAI_API_KEY")
    model = (os.getenv("DASHSCOPE_MODEL") or "qwen3-max").strip()
    if not api_key:
        raise RuntimeError("未配置大模型：请设置 DASHSCOPE_API_KEY（或 OPENAI_API_KEY）。")
    max_tokens = int(os.getenv("DASHSCOPE_WORD_MAX_TOKENS", "8192"))
    timeout_sec = int(os.getenv("DASHSCOPE_WORD_TIMEOUT_SEC", "180"))
    profile = student_record.get("profile") if isinstance(student_record.get("profile"), dict) else {}
    system_prompt = (
        "你是学业报告文书填充助手。任务：根据学生数据，逐个判断 Word 表格单元格应填写的最终文本。"
        "你必须保留非目标单元格原文；仅在明确可映射到学生数据时改写。"
        "若单元格原文非空，final_text 不得为空字符串（禁止误清空）。"
        "若单元格含 {{key}} 占位符，优先按 key 替换。"
        "若单元格是“字段名:原值/字段名：原值”，可根据字段名替换。"
        "禁止改写无占位符、无「标签:」结构且仅为短固定词的单元格（例如单独一行人名、课程名）；此类必须原样返回 original_text。"
        "表格同一行若存在左右两列（table_cells 中相同 table_index、row_index 下 col_index 为 0 与 1）："
        "左格只保留「标签：」标签本身或单独标题，具体数字与正文必须写在右格（col_index 较大的单元格）；"
        "禁止把「课程次数：1」这类「标签+值」整段只写在左格而右格为空。"
        "转案日期、专业、委托产品、对接顾问等：若 value_map 无键名，可从 profile.basic / 全量 student_record 按常见语义推断并填写；勿留空。"
        "若 table_cells[].suggest_infer 为 true：表示该格为空、仅占位或仅有「标签：」无值，你必须结合 student_record / profile / value_map 推理可填内容；"
        "能从其他字段合理推出则填写具体文字，确实无依据时可写「—」或一两字说明，不要留空占位。"
        "不要输出解释，只输出 JSON。"
    )
    payload = {
        "student_record": student_record,
        "profile": profile,
        "value_map": values,
        "label_aliases": aliases,
        "table_cells": cells,
    }
    user_prompt = (
        "请输出如下 JSON：\n"
        "{\n"
        '  "cells": [\n'
        '    {"cell_id":"t0_r0_c0","final_text":"..."}\n'
        "  ]\n"
        "}\n"
        "要求：\n"
        "1) 只输出存在于输入中的 cell_id；\n"
        "2) 每个 cell_id 仅出现一次；\n"
        "3) final_text 为最终写回文本；\n"
        "4) 若不应修改，final_text 返回 original_text；\n"
        "5) 全文保持中文风格，不新增无关内容；\n"
        "6) 对无冒号、无{{}}的短行固定文案（如模板里已写好的教服老师姓名）必须保持原样；\n"
        "7) 双列表格同行左右两格分工：左标签、右数值或说明；\n"
        "8) 尽量补全转案日期、专业、委托产品、对接顾问等有迹可循的字段；\n"
        "9) 凡 suggest_infer 为 true 的单元格必须输出推理后的 final_text（不得与原文一样是空或仅「/」）。\n\n"
        f"输入数据：{json.dumps(payload, ensure_ascii=False)}"
    )
    response = requests.post(
        "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        json={
            "model": model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            "temperature": 0.15,
            "max_tokens": max_tokens,
        },
        timeout=timeout_sec,
    )
    if response.status_code >= 400:
        raise RuntimeError(f"Word 生成模型调用失败：HTTP {response.status_code}")
    data = response.json()
    try:
        content = data["choices"][0]["message"]["content"]
    except (KeyError, IndexError, TypeError) as exc:
        raise RuntimeError("Word 生成模型返回格式异常。") from exc
    try:
        parsed = extract_json_from_text(content)
    except json.JSONDecodeError as exc:
        raise RuntimeError("Word 生成模型未返回合法 JSON。") from exc
    rows = parsed.get("cells") if isinstance(parsed, dict) else None
    if not isinstance(rows, list):
        raise RuntimeError("Word 生成模型返回缺少 cells 列表。")
    out: dict[str, str] = {}
    for row in rows:
        if not isinstance(row, dict):
            continue
        cid = str(row.get("cell_id") or "").strip()
        if not cid:
            continue
        out[cid] = str(row.get("final_text") or "")
    return out


def fill_word_table_for_student(
    task_id: str,
    student_data_id: str,
    *,
    output_path: Path | None = None,
) -> dict[str, Any]:
    tid = (task_id or "").strip()
    sid = (student_data_id or "").strip()
    if not is_safe_task_id(tid):
        raise ValueError("无效的模板任务 ID。")
    if not sid:
        raise ValueError("缺少学生数据 ID。")
    template = config.UPLOAD_DIR / f"{tid}.docx"
    if not template.is_file():
        raise ValueError("未找到 Word 模板文件，请先上传 .docx。")
    rec = db.get_student_record(sid)
    if not rec:
        raise ValueError("未找到学生数据，请重新选择。")
    values, aliases = _build_student_value_map(rec)
    if not values:
        raise ValueError("学生数据为空，无法回填。")

    doc = Document(str(template))
    cells_meta, table_count = _extract_table_cells(doc)
    if table_count <= 0:
        raise ValueError("Word 模板中未检测到表格，无法执行表格回填。")

    deterministic_cells = 0
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                orig = cell.text or ""
                new_t = _deterministic_cell_text(orig, values, aliases)
                if new_t != orig:
                    set_table_cell_text_preserve_style(cell, new_t)
                    deterministic_cells += 1

    two_column_row_fixes = _fix_two_column_rows_until_stable(doc, values, aliases)

    cells_meta, table_count = _extract_table_cells(doc)

    llm_updates = _llm_fill_table_cells(
        student_record=rec,
        values=values,
        aliases=aliases,
        cells=cells_meta,
    )

    touched_cells = 0
    placeholder_hits = 0
    kv_rewrite_hits = 0
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cid = f"t{t_idx}_r{r_idx}_c{c_idx}"
                original = cell.text or ""
                updated = llm_updates.get(cid, original)
                # 禁止用空串覆盖原有非空内容，避免模型误返回 final_text="" 导致整格被清空
                if (original or "").strip() and not str(updated or "").strip():
                    updated = original
                # 禁止模型用另一段非空文字覆盖「短固定模板」（无占位符、非 KV 首行、非长文）
                if (
                    str(updated or "").strip()
                    and str(updated or "").strip() != str(original or "").strip()
                    and not _should_allow_llm_cell_update(original)
                ):
                    updated = original
                if updated != original:
                    touched_cells += 1
                    placeholder_hits += len(_PLACEHOLDER_RE.findall(original))
                    kv_rewrite_hits += 1 if _KV_LINE_RE.match(original or "") else 0
                    set_table_cell_text_preserve_style(cell, updated)

    two_column_row_fixes += _fix_two_column_rows_until_stable(doc, values, aliases)

    if output_path is None:
        out = config.FILLED_EXPORT_DIR / f"{tid}_filled.docx"
    else:
        out = output_path
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return {
        "ok": True,
        "task_id": tid,
        "student_data_id": sid,
        "output_kind": "docx",
        "table_count": table_count,
        "deterministic_cells": deterministic_cells,
        "touched_cells": touched_cells,
        "placeholder_hits": placeholder_hits,
        "kv_rewrite_hits": kv_rewrite_hits,
        "two_column_row_fixes": two_column_row_fixes,
        "download_path": str(out),
    }

